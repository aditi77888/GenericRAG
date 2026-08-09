from pathlib import Path

from docx import Document as DocxDocument

from loaders.base_loader import BaseLoader
from models.document import Document


class DOCXLoader(BaseLoader):

    def load(self, source):

        path = self._validate_file(source)

        text = self._read_docx(path)

        document = self._create_document(
            text,
            path
        )

        return [document]

    def _validate_file(self, source):

        path = Path(source)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        if path.suffix.lower() != ".docx":
            raise ValueError("Only .docx files are supported.")

        return path

    def _read_docx(self, path):

        doc = DocxDocument(path)

        paragraphs = []

        for paragraph in doc.paragraphs:

            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        return "\n".join(paragraphs)

    def _create_document(
            self,
            text,
            path
    ) -> Document:

        metadata = {
            "source": str(path),
            "file_name": path.name,
            "file_type": "docx"
        }

        return Document(
            content=text,
            metadata=metadata
        )